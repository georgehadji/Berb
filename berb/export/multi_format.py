"""Multi-format export engine for research papers.

This module provides unified export supporting all output formats:
- PDF (compiled LaTeX)
- LaTeX source
- Word (.docx)
- Markdown
- HTML
- BibTeX
- CSL-JSON
- Overleaf ZIP

# Author: Georgios-Chrysovalantis Chatzivantsidis
"""

from __future__ import annotations

import json
import logging
from enum import Enum
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    """Supported export formats.

    Attributes:
        PDF: Compiled PDF document
        LATEX: LaTeX source files
        WORD: Microsoft Word (.docx)
        MARKDOWN: Markdown (.md)
        HTML: Web-readable HTML
        BIBTEX: BibTeX bibliography only
        CSL_JSON: CSL-JSON for reference managers
        OVERLEAF_ZIP: Overleaf-ready ZIP package
    """

    PDF = "pdf"
    LATEX = "latex"
    WORD = "word"
    MARKDOWN = "markdown"
    HTML = "html"
    BIBTEX = "bibtex"
    CSL_JSON = "csl-json"
    OVERLEAF_ZIP = "overleaf-zip"


class ExportConfig(BaseModel):
    """Export configuration.

    Attributes:
        formats: List of formats to export
        output_dir: Output directory
        template: LaTeX template name
        include_sections: Which sections to include
        include_bibliography: Include bibliography
        include_figures: Include figures
    """

    formats: list[ExportFormat] = Field(
        default_factory=lambda: [ExportFormat.LATEX, ExportFormat.MARKDOWN]
    )
    output_dir: str = "output"
    template: str = "article"
    include_sections: list[str] = Field(
        default_factory=lambda: [
            "abstract",
            "introduction",
            "methods",
            "results",
            "discussion",
            "conclusion",
        ]
    )
    include_bibliography: bool = True
    include_figures: bool = True


class ExportResult(BaseModel):
    """Result of multi-format export.

    Attributes:
        exported_formats: List of successfully exported formats
        output_paths: Paths to exported files
        errors: Any errors encountered
        total_size_bytes: Total size of all exports
    """

    exported_formats: list[str] = Field(default_factory=list)
    output_paths: dict[str, str] = Field(default_factory=dict)
    errors: list[str] = Field(default_factory=list)
    total_size_bytes: int = 0


class MultiFormatExporter:
    """Export papers in multiple formats simultaneously.

    This exporter provides a unified interface for exporting
    research papers in various formats for different use cases.

    Usage::

        exporter = MultiFormatExporter(
            config=ExportConfig(
                formats=[
                    ExportFormat.LATEX,
                    ExportFormat.WORD,
                    ExportFormat.HTML,
                ],
            ),
        )

        result = await exporter.export(
            paper=paper_content,
            references=references,
            figures=figures,
        )

    Attributes:
        config: Export configuration
    """

    def __init__(self, config: ExportConfig | None = None):
        """Initialize multi-format exporter.

        Args:
            config: Export configuration
        """
        self.config = config or ExportConfig()

        # Import LaTeX exporter lazily
        self._latex_exporter = None

    async def export(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None = None,
    ) -> ExportResult:
        """Export paper in multiple formats.

        Args:
            paper: Paper content and metadata
            references: List of references
            figures: List of figure file paths

        Returns:
            ExportResult with all exported files
        """
        result = ExportResult()
        output_dir = Path(self.config.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        for fmt in self.config.formats:
            try:
                if fmt == ExportFormat.LATEX:
                    path = await self._export_latex(
                        paper, references, figures, output_dir
                    )
                    result.exported_formats.append("latex")
                    result.output_paths["latex"] = str(path)

                elif fmt == ExportFormat.MARKDOWN:
                    path = self._export_markdown(paper, references, output_dir)
                    result.exported_formats.append("markdown")
                    result.output_paths["markdown"] = str(path)

                elif fmt == ExportFormat.HTML:
                    path = await self._export_html(
                        paper, references, figures, output_dir
                    )
                    result.exported_formats.append("html")
                    result.output_paths["html"] = str(path)

                elif fmt == ExportFormat.WORD:
                    path = await self._export_word(
                        paper, references, figures, output_dir
                    )
                    result.exported_formats.append("word")
                    result.output_paths["word"] = str(path)

                elif fmt == ExportFormat.BIBTEX:
                    path = self._export_bibtex(references, output_dir)
                    result.exported_formats.append("bibtex")
                    result.output_paths["bibtex"] = str(path)

                elif fmt == ExportFormat.CSL_JSON:
                    path = self._export_csl_json(references, output_dir)
                    result.exported_formats.append("csl-json")
                    result.output_paths["csl-json"] = str(path)

                elif fmt == ExportFormat.OVERLEAF_ZIP:
                    path = await self._export_overleaf(
                        paper, references, figures, output_dir
                    )
                    result.exported_formats.append("overleaf-zip")
                    result.output_paths["overleaf-zip"] = str(path)

                elif fmt == ExportFormat.PDF:
                    # PDF requires LaTeX compilation
                    path = await self._export_pdf(
                        paper, references, figures, output_dir
                    )
                    result.exported_formats.append("pdf")
                    result.output_paths["pdf"] = str(path)

            except Exception as e:
                logger.error(f"Failed to export {fmt.value}: {e}")
                result.errors.append(f"{fmt.value}: {str(e)}")

        # Calculate total size
        result.total_size_bytes = self._calculate_total_size(result.output_paths)

        logger.info(
            f"Export complete: {len(result.exported_formats)} formats, "
            f"{result.total_size_bytes / 1024:.1f} KB total"
        )

        return result

    async def _export_latex(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None,
        output_dir: Path,
    ) -> Path:
        """Export to LaTeX.

        Args:
            paper: Paper content
            references: References
            figures: Figure paths
            output_dir: Output directory

        Returns:
            Path to main.tex
        """
        from berb.export.latex_exporter import LaTeXExporter, LaTeXExporterConfig

        exporter = LaTeXExporter(
            config=LaTeXExporterConfig(template=self.config.template),
        )

        latex_dir = output_dir / "latex"
        project = await exporter.export_latex_project(
            paper=paper,
            references=references,
            output_dir=latex_dir,
            figures=figures,
        )

        return latex_dir / "main.tex"

    def _export_markdown(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        output_dir: Path,
    ) -> Path:
        """Export to Markdown.

        Args:
            paper: Paper content
            references: References
            output_dir: Output directory

        Returns:
            Path to .md file
        """
        lines = []

        # Title and authors
        title = paper.get("title", "Untitled")
        authors = paper.get("authors", [])
        abstract = paper.get("abstract", "")

        lines.append(f"# {title}\n")
        lines.append(f"**Authors:** {', '.join(authors)}\n")
        lines.append(f"\n## Abstract\n\n{abstract}\n")

        # Sections
        section_order = [
            ("introduction", "Introduction"),
            ("related_work", "Related Work"),
            ("methods", "Methods"),
            ("results", "Results"),
            ("discussion", "Discussion"),
            ("conclusion", "Conclusion"),
        ]

        for key, title in section_order:
            content = paper.get(key, "")
            if content:
                lines.append(f"\n## {title}\n\n{content}\n")

        # Bibliography
        if self.config.include_bibliography:
            lines.append("\n## References\n\n")
            for i, ref in enumerate(references, 1):
                ref_str = self._format_reference_markdown(ref)
                lines.append(f"{i}. {ref_str}\n")

        # Write file
        output_path = output_dir / "paper.md"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

        return output_path

    def _format_reference_markdown(
        self,
        ref: dict[str, Any],
    ) -> str:
        """Format reference as Markdown.

        Args:
            ref: Reference metadata

        Returns:
            Formatted reference string
        """
        authors = ref.get("authors", [])
        year = ref.get("year", "n.d.")
        title = ref.get("title", "Untitled")
        venue = ref.get("venue", "")

        author_str = " and ".join(authors) if authors else "Anonymous"

        ref_str = f"{author_str} ({year}). {title}."
        if venue:
            ref_str += f" *{venue}*."

        if ref.get("doi"):
            ref_str += f" https://doi.org/{ref['doi']}"

        return ref_str

    async def _export_html(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None,
        output_dir: Path,
    ) -> Path:
        """Export to HTML.

        Args:
            paper: Paper content
            references: References
            figures: Figure paths
            output_dir: Output directory

        Returns:
            Path to .html file
        """
        title = paper.get("title", "Untitled")
        authors = paper.get("authors", [])

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; border-bottom: 1px solid #ddd; padding-bottom: 5px; }}
        .authors {{ color: #666; font-style: italic; }}
        .abstract {{ background: #f5f5f5; padding: 15px; border-left: 3px solid #333; }}
        .reference {{ margin-bottom: 10px; }}
        figure {{ margin: 20px 0; }}
        figcaption {{ font-style: italic; color: #666; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="authors">{', '.join(authors)}</p>
    
    <div class="abstract">
        <h2>Abstract</h2>
        <p>{paper.get('abstract', '')}</p>
    </div>
"""

        # Add sections
        sections = [
            ("introduction", "Introduction"),
            ("methods", "Methods"),
            ("results", "Results"),
            ("discussion", "Discussion"),
            ("conclusion", "Conclusion"),
        ]

        for key, section_title in sections:
            content = paper.get(key, "")
            if content:
                html += f"    <h2>{section_title}</h2>\n    {content}\n"

        # Add figures
        if self.config.include_figures and figures:
            html += "    <h2>Figures</h2>\n"
            for i, fig_path in enumerate(figures, 1):
                html += f'    <figure><img src="{fig_path}" alt="Figure {i}"><figcaption>Figure {i}</figcaption></figure>\n'

        # Add references
        if self.config.include_bibliography:
            html += "    <h2>References</h2>\n"
            for i, ref in enumerate(references, 1):
                ref_str = self._format_reference_markdown(ref)
                html += f'    <div class="reference">{i}. {ref_str}</div>\n'

        html += """
</body>
</html>
"""

        output_path = output_dir / "paper.html"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)

        return output_path

    async def _export_word(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None,
        output_dir: Path,
    ) -> Path:
        """Export to Word (.docx).

        Args:
            paper: Paper content
            references: References
            figures: Figure paths
            output_dir: Output directory

        Returns:
            Path to .docx file
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            # python-docx not installed - fall back to Markdown
            logger.warning("python-docx not installed, exporting as Markdown instead")
            return self._export_markdown(paper, references, output_dir)

        doc = Document()

        # Title
        title = paper.get("title", "Untitled")
        doc.add_heading(title, 0)

        # Authors
        authors = paper.get("authors", [])
        doc.add_paragraph(", ".join(authors), style="Intense Quote")

        # Abstract
        abstract = paper.get("abstract", "")
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract)

        # Sections
        sections = [
            ("introduction", "Introduction"),
            ("related_work", "Related Work"),
            ("methods", "Methods"),
            ("results", "Results"),
            ("discussion", "Discussion"),
            ("conclusion", "Conclusion"),
        ]

        for key, section_title in sections:
            content = paper.get(key, "")
            if content:
                doc.add_heading(section_title, level=1)
                doc.add_paragraph(content)

        # Figures
        if self.config.include_figures and figures:
            doc.add_heading("Figures", level=1)
            for i, fig_path in enumerate(figures, 1):
                try:
                    doc.add_picture(fig_path, width=Inches(6))
                    doc.add_paragraph(f"Figure {i}", style="Caption")
                except Exception as e:
                    logger.warning(f"Failed to add figure {fig_path}: {e}")

        # References
        if self.config.include_bibliography:
            doc.add_heading("References", level=1)
            for i, ref in enumerate(references, 1):
                ref_str = self._format_reference_markdown(ref)
                doc.add_paragraph(f"{i}. {ref_str}")

        output_path = output_dir / "paper.docx"
        doc.save(output_path)

        return output_path

    def _export_bibtex(
        self,
        references: list[dict[str, Any]],
        output_dir: Path,
    ) -> Path:
        """Export to BibTeX.

        Args:
            references: References
            output_dir: Output directory

        Returns:
            Path to .bib file
        """
        from berb.writing.citation_styles import CitationFormatter

        formatter = CitationFormatter()
        bibtex = formatter.to_bibtex(references)

        output_path = output_dir / "references.bib"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(bibtex)

        return output_path

    def _export_csl_json(
        self,
        references: list[dict[str, Any]],
        output_dir: Path,
    ) -> Path:
        """Export to CSL-JSON.

        Args:
            references: References
            output_dir: Output directory

        Returns:
            Path to .json file
        """
        csl_items = []

        for ref in references:
            # Map to CSL-JSON format
            item_type = self._map_to_csl_type(ref.get("type", "article"))

            csl_item = {
                "type": item_type,
                "id": ref.get("doi", ref.get("id", "")),
                "title": ref.get("title", ""),
                "author": [
                    {"family": a.split()[-1], "given": " ".join(a.split()[:-1])}
                    for a in ref.get("authors", [])
                ],
                "issued": {"date-parts": [[ref.get("year", "")]]},
            }

            if ref.get("venue"):
                csl_item["container-title"] = ref.get("venue")

            if ref.get("volume"):
                csl_item["volume"] = str(ref.get("volume"))

            if ref.get("pages"):
                csl_item["page"] = str(ref.get("pages"))

            if ref.get("doi"):
                csl_item["DOI"] = ref.get("doi")

            csl_items.append(csl_item)

        output_path = output_dir / "references.json"
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(csl_items, f, indent=2)

        return output_path

    def _map_to_csl_type(
        self,
        ref_type: str,
    ) -> str:
        """Map reference type to CSL type.

        Args:
            ref_type: Reference type

        Returns:
            CSL type string
        """
        type_mapping = {
            "article": "article-journal",
            "inproceedings": "paper-conference",
            "book": "book",
            "incollection": "chapter",
            "phdthesis": "thesis",
            "mastersthesis": "thesis",
            "techreport": "report",
            "manual": "manual",
            "misc": "webpage",
        }

        return type_mapping.get(ref_type, "article-journal")

    async def _export_overleaf(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None,
        output_dir: Path,
    ) -> Path:
        """Export to Overleaf ZIP.

        Args:
            paper: Paper content
            references: References
            figures: Figure paths
            output_dir: Output directory

        Returns:
            Path to .zip file
        """
        from berb.export.latex_exporter import LaTeXExporter, LaTeXExporterConfig

        exporter = LaTeXExporter(
            config=LaTeXExporterConfig(template=self.config.template),
        )

        latex_dir = output_dir / "latex_temp"
        project = await exporter.export_latex_project(
            paper=paper,
            references=references,
            output_dir=latex_dir,
            figures=figures,
        )

        zip_path = await exporter.export_overleaf_zip(
            project=project,
            output_path=output_dir / "overleaf_package.zip",
        )

        # Clean up temp directory
        import shutil
        shutil.rmtree(latex_dir, ignore_errors=True)

        return zip_path

    async def _export_pdf(
        self,
        paper: dict[str, Any],
        references: list[dict[str, Any]],
        figures: list[str] | None,
        output_dir: Path,
    ) -> Path | None:
        """Export to PDF (requires LaTeX installation).

        Args:
            paper: Paper content
            references: References
            figures: Figure paths
            output_dir: Output directory

        Returns:
            Path to .pdf file or None if compilation fails
        """
        import subprocess

        # First export LaTeX
        latex_path = await self._export_latex(
            paper, references, figures, output_dir
        )

        latex_dir = latex_path.parent

        try:
            # Run pdflatex
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "main.tex"],
                cwd=latex_dir,
                capture_output=True,
                timeout=120,
            )

            # Run bibtex
            subprocess.run(
                ["bibtex", "main"],
                cwd=latex_dir,
                capture_output=True,
                timeout=60,
            )

            # Run pdflatex twice more for references
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "main.tex"],
                cwd=latex_dir,
                capture_output=True,
                timeout=120,
            )
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "main.tex"],
                cwd=latex_dir,
                capture_output=True,
                timeout=120,
            )

            pdf_path = latex_dir / "main.pdf"
            if pdf_path.exists():
                return pdf_path
            else:
                logger.warning("PDF compilation failed - no output PDF")
                return None

        except Exception as e:
            logger.error(f"PDF compilation failed: {e}")
            return None

    def _calculate_total_size(
        self,
        output_paths: dict[str, str],
    ) -> int:
        """Calculate total size of exported files.

        Args:
            output_paths: Dictionary of format -> path

        Returns:
            Total size in bytes
        """
        total = 0
        for path_str in output_paths.values():
            path = Path(path_str)
            if path.exists():
                total += path.stat().st_size
            elif path.is_dir():
                # For directories (like LaTeX export), sum all files
                for file in path.rglob("*"):
                    if file.is_file():
                        total += file.stat().st_size
        return total


# Convenience function
async def export_paper(
    paper: dict[str, Any],
    references: list[dict[str, Any]],
    formats: list[ExportFormat] | None = None,
    output_dir: str = "output",
    template: str = "article",
) -> ExportResult:
    """Convenience function for multi-format export.

    Args:
        paper: Paper content
        references: References
        formats: Formats to export
        output_dir: Output directory
        template: LaTeX template

    Returns:
        ExportResult
    """
    config = ExportConfig(
        formats=formats or [ExportFormat.LATEX, ExportFormat.MARKDOWN, ExportFormat.HTML],
        output_dir=output_dir,
        template=template,
    )

    exporter = MultiFormatExporter(config=config)
    return await exporter.export(paper, references)
